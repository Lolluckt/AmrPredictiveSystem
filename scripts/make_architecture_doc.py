"""Generate D:\\диплом\\Архітектура.docx — an architectural description
of the project (Ukrainian, with ASCII schemes).

Run once:
    python scripts/make_architecture_doc.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parent.parent / "Архітектура.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def add_table(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light List Accent 1"
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"
    for r, row in enumerate(rows, start=1):
        for i, text in enumerate(row):
            cell = table.rows[r].cells[i]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, width_cm in enumerate(col_widths):
                row.cells[i].width = Cm(width_cm)
    return table


def build() -> None:
    doc = Document()

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)


    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Архітектура інформаційної системи\n"
        "збору телеметрії та предиктивного обслуговування АМР"
    )
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Пояснювальна записка до проєктної частини дипломної роботи")
    r.italic = True
    r.font.size = Pt(12)

    doc.add_paragraph()
    tl = doc.add_paragraph()
    tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tl.add_run("Тема: «Інформаційна система збору телеметрії та предиктивного "
               "обслуговування автономної наземної колісної платформи в умовах "
               "цехової виробничої лінії з використанням технологій цифрових двійників»").italic = True

    doc.add_page_break()


    add_heading(doc, "1. Загальний огляд системи", 1)
    add_paragraph(doc,
        "Система реалізує повний життєвий цикл даних — від генерації на "
        "цифровому двійнику цеху у Webots, через транспорт MQTT і зберігання у "
        "PostgreSQL, і до візуалізації та роботи з тікетами у веб-інтерфейсі "
        "React. Три основні програмні компоненти (симулятор, бекенд, фронтенд) "
        "комунікують через два чітко визначені протоколи: MQTT для потокової "
        "телеметрії та HTTP/JSON з JWT-авторизацією для REST-API.")

    add_paragraph(doc, "Загальна логічна схема:", bold=True)
    add_code(doc,
        "┌───────────────────────────────────────────────────────────────┐\n"
        "│                  Webots (цифровий двійник цеху)               │\n"
        "│   ┌──────────┐   ┌──────────┐   ┌──────────┐                  │\n"
        "│   │ AMR-01   │   │ AMR-02   │   │ AMR-03   │   Python         │\n"
        "│   │ (Python) │   │ (Python) │   │ (Python) │   amr_controller │\n"
        "│   └────┬─────┘   └────┬─────┘   └────┬─────┘                  │\n"
        "└────────┼──────────────┼───────────────┼────────────────────────┘\n"
        "         └──────────────┼───────────────┘\n"
        "                  MQTT publish\n"
        "         factory/line_1/amr_xx/telemetry/{section}\n"
        "                  ▼\n"
        "         ┌────────────────────┐\n"
        "         │ Mosquitto (broker) │\n"
        "         └──────────┬─────────┘\n"
        "                    │ subscribe  factory/+/+/telemetry/#\n"
        "                    ▼\n"
        "         ┌─────────────────────────────────────────────────────┐\n"
        "         │                    Backend (FastAPI)                │\n"
        "         │  ┌────────────┐  ┌─────────────┐  ┌────────────┐    │\n"
        "         │  │ MQTT ingest│→ │ Anomaly     │→ │ Predictive │    │\n"
        "         │  │  + queue   │  │ detection   │  │ RUL + health│   │\n"
        "         │  └─────┬──────┘  └─────┬───────┘  └─────┬──────┘    │\n"
        "         │        ▼               ▼                ▼           │\n"
        "         │      ┌───────────────────────────────────────┐      │\n"
        "         │      │           PostgreSQL 16               │      │\n"
        "         │      │ users · robots · telemetry_snapshots  │      │\n"
        "         │      │ alert_rules · anomaly_events          │      │\n"
        "         │      │ tickets · missions · audit_log        │      │\n"
        "         │      └───────────────────────────────────────┘      │\n"
        "         │        ▲                                             │\n"
        "         │        │ REST (FastAPI routers + JWT-authz)         │\n"
        "         └────────┼─────────────────────────────────────────────┘\n"
        "                  ▼\n"
        "         ┌─────────────────────────────────────────────────────┐\n"
        "         │             Frontend (React + TS + Tailwind)        │\n"
        "         │  Dashboard · Robots · Predictive · Alerts ·         │\n"
        "         │  Tickets (CMMS Kanban) · Missions · Admin/Users     │\n"
        "         └─────────────────────────────────────────────────────┘")


    add_heading(doc, "2. Структура проєкту", 1)
    add_paragraph(doc, "Кореневий каталог диплому:")
    add_code(doc,
        "диплом/\n"
        "├── webots_project/        цифровий двійник (Webots R2025a)\n"
        "│   ├── worlds/factory_floor.wbt       сцена цеху 40×24 м\n"
        "│   ├── controllers/amr_controller/    Python-контролер AMR\n"
        "│   └── README.md\n"
        "├── backend/                FastAPI + SQLAlchemy (async)\n"
        "│   ├── app/\n"
        "│   │   ├── main.py         FastAPI app-factory + lifespan\n"
        "│   │   ├── config.py       pydantic-settings\n"
        "│   │   ├── db.py           async engine + session\n"
        "│   │   ├── core/           security · roles · deps (RBAC)\n"
        "│   │   ├── models/         SQLAlchemy ORM (8 модулів)\n"
        "│   │   ├── schemas/        Pydantic request/response\n"
        "│   │   ├── routers/        auth · users · robots · telemetry\n"
        "│   │   │                   · alerts · tickets · missions · predictive\n"
        "│   │   ├── services/       mqtt_ingest · mqtt_publisher\n"
        "│   │   │                   · anomaly · predictive (RUL)\n"
        "│   │   └── seed.py         bootstrap даних (idempotent)\n"
        "│   ├── Dockerfile\n"
        "│   └── requirements.txt\n"
        "├── frontend/               React + Vite + TS + Tailwind\n"
        "│   └── src/\n"
        "│       ├── api/            axios клієнт + endpoints\n"
        "│       ├── components/     Layout, PageHeader, Badge, guards\n"
        "│       ├── pages/          Login · Dashboard · Robots · RobotDetail\n"
        "│       │                   · Predictive · Alerts · Tickets\n"
        "│       │                   · Missions · Users\n"
        "│       ├── store/auth.ts   Zustand store (persisted)\n"
        "│       ├── types.ts\n"
        "│       └── App.tsx         role-aware routing\n"
        "├── db.txt                  довідкова DDL (50+ таблиць)\n"
        "├── diplom_er_diagrams.html ER-діаграми предметної області\n"
        "├── docker-compose.yml      postgres + mosquitto + backend + frontend\n"
        "└── Архітектура.docx        (цей документ)")


    add_heading(doc, "3. Цифровий двійник у Webots", 1)
    add_heading(doc, "3.1. Сцена цеху", 2)
    add_paragraph(doc,
        "Сцена factory_floor.wbt реалізує цех 40 × 24 м з 9-ма функціональними "
        "зонами, широкими коридорами (3 м вертикальний, 2 м поперечний), "
        "освітленням, жовтими маркерами безпеки, HMI-панелями та повним набором "
        "обладнання (стелажі, столи збирання, зварювальні пости, конвеєр "
        "пакування, 3 зарядні станції, столи контролю якості).")

    add_paragraph(doc, "Компонування зон (вигляд зверху):", bold=True)
    add_code(doc,
        "  y=24 ┌──────────┬──────────────────────────┬──────────┬──────────┐\n"
        "       │          │                          │          │          │\n"
        "       │   QC     │    Assembly B            │ Packaging│ Finished │\n"
        "       │          │                          │          │ warehouse│\n"
        "  y=15 ├──────────┤                          ├──────────┤          │\n"
        "       │ Charging │                          │          │          │\n"
        "  y=11 ├──────────┴── CORRIDOR MID (2 m) ────┴──────────┤          │\n"
        "  y=9  ├──────────┬──────────────────────────┬──────────┤          │\n"
        "       │ Warehouse│    Assembly A            │ Welding  │          │\n"
        "       │   raw    │                          │          │          │\n"
        "  y=0  └──────────┴──────────────────────────┴──────────┴──────────┘\n"
        "       x=0       x=10/13                    x=23       x=32     x=40\n"
        "                  └── corridor main (3 m) ──┘")

    add_heading(doc, "3.2. Конструкція АМР", 2)
    add_paragraph(doc,
        "Кожен робот — циліндрична платформа Ø 0.30 м, маса 12 кг, диференційний "
        "привід (два активних колеса Ø 0.12 м + два пасивних керамічних "
        "кастери). Низький центр мас стабілізує рух на високих курсових "
        "швидкостях.")
    add_paragraph(doc, "Сенсори (однакові у всіх трьох роботів):")
    for b in [
        "2× RotationalMotor + PositionSensor (енкодери коліс, 1024 PPR)",
        "GPS для абсолютної позиції (ENU)",
        "Compass для heading",
        "InertialUnit (roll/pitch/yaw) + Accelerometer + Gyro",
        "LiDAR 360°, 32 промені, дальність 5 м, шум 5 мм",
        "8-секторний сонарний ring (радіальне покриття кожні 45°)",
        "LED статусу (0 — зелений/idle, 1 — жовтий/working, 2 — червоний/low battery, 3 — синій/charging)",
    ]: add_bullet(doc, b)

    add_heading(doc, "3.3. Алгоритм руху", 2)
    add_paragraph(doc,
        "Попередня версія контролера рухалась лише за принципом prop-heading, "
        "через що робот застрягав у кутах стелажів. Нова реалізація має три рівні:")
    add_bullet(doc,
        "Граф вузлів: 19 іменованих waypoints, розташованих у коридорах та "
        "на drop-off маркерах. Ребра графа — лише через відкриті коридори.")
    add_bullet(doc,
        "BFS-пошук шляху: при зміні цілі місії контролер знаходить "
        "найкоротший шлях у графі й слідує ним waypoint за waypoint.")
    add_bullet(doc,
        "Каскадне керування: P-регулятор курсу → трансформація у швидкості "
        "колес + гальмування за відстанню до цільової точки. Злиття LiDAR + "
        "сонарів дає 3-рівневе уникнення перешкод: slow-bias → hard swerve → "
        "stop-and-turn.")
    add_bullet(doc,
        "Stuck-recovery: детекція < 8 см зміщення за 2.5 с → реверс 1.5 с із "
        "міняною стороною; після 4 спроб waypoint пропускається.")

    add_heading(doc, "3.4. Імітація деградації компонентів", 2)
    add_paragraph(doc,
        "Контролер містить модель Degradation, яка тік за тіком оновлює стан "
        "усіх компонентів з фізично вмотивованими залежностями. "
        "Це дозволяє бекенду формувати реалістичну аналітику без ML:")
    add_table(doc,
        ["Компонент", "Параметри", "Формула"],
        [
            ["Батарея", "SoC, SoH, V, I, T, R_internal",
             "ΔSoC ∝ I·dt; ΔSoH ∝ load·dt; V_sag ∝ fade·load"],
            ["Двигун",  "T_motor, вібрація, ККД, bearing_wear",
             "ΔT ∝ (v/v_max)²·dt; vib = 0.1 + wear·4 + shocks"],
            ["Енкодер", "slip (rad)",
             "ΔΦ ∝ 𝒩(0,0.002)·dt (при fault_encoder_drift)"],
            ["Привід (brake)", "віртуальний коефіцієнт крутного моменту",
             "motor_output *= 0.3 (якщо fault_brake_stuck)"],
        ],
        col_widths=[3.0, 5.0, 8.5])
    add_paragraph(doc,
        "Вбудовано п’ять режимів інжекції несправностей (bearing_right, "
        "thermal_left, battery_fade, encoder_drift, brake_stuck), які бекенд "
        "може вмикати через MQTT-команду «inject_fault». Це — ключ до "
        "демонстрації предиктивного обслуговування в живому режимі.")

    doc.add_page_break()


    add_heading(doc, "4. Бекенд (FastAPI)", 1)
    add_heading(doc, "4.1. Стек і структура", 2)
    add_paragraph(doc, "Python 3.12 · FastAPI 0.115 · SQLAlchemy 2.0 (async) · "
                        "asyncpg · Pydantic v2 · paho-mqtt · python-jose (JWT).")
    add_code(doc,
        "backend/app/\n"
        "├── main.py\n"
        "├── config.py        (pydantic-settings, .env)\n"
        "├── db.py            (async engine, session_scope)\n"
        "├── core/\n"
        "│   ├── roles.py     (enum Role: 3 ролі + таблиця Permission)\n"
        "│   ├── security.py  (bcrypt + JWT access/refresh)\n"
        "│   └── deps.py      (current_user, require_permission)\n"
        "├── models/          (ORM — 8 доменів)\n"
        "├── schemas/         (Pydantic DTO)\n"
        "├── routers/\n"
        "│   ├── auth.py         /api/auth/{login,refresh,logout,me}\n"
        "│   ├── users.py        /api/users*                       (admin)\n"
        "│   ├── robots.py       /api/robots{,/{id},/{id}/command}\n"
        "│   ├── telemetry.py    /api/telemetry/{id}/{latest,history,series}\n"
        "│   ├── alerts.py       /api/{alert-rules,anomalies}*\n"
        "│   ├── tickets.py      /api/tickets*                     (CMMS)\n"
        "│   ├── missions.py     /api/missions*\n"
        "│   └── predictive.py   /api/predictive/{fleet,robots/{id}/{health,rul}}\n"
        "└── services/\n"
        "    ├── mqtt_ingest.py     subscribe → queue → DB + anomaly\n"
        "    ├── mqtt_publisher.py  outbound команди (stop/inject_fault)\n"
        "    ├── anomaly.py         eval alert_rules проти snapshot\n"
        "    └── predictive.py      компонент health + RUL")

    add_heading(doc, "4.2. Модель даних (ORM)", 2)
    add_paragraph(doc,
        "Реалізовано 15 ключових таблиць із повної 50+-табличної схеми "
        "(db.txt). Вибрано підмножину, що відповідає функціоналу 3 ролей і "
        "усуває дублювання без втрати демонстрації PdM-циклу:")
    add_table(doc,
        ["Домен", "Таблиці", "Призначення"],
        [
            ["Користувачі", "users, refresh_tokens", "Автентифікація, 3 ролі"],
            ["Фабрика",     "factories, production_lines, workshop_zones, charging_stations",
             "Ієрархія підприємства"],
            ["Флот",        "robots, robot_components", "Реєстр AMR, інвентар деталей"],
            ["Телеметрія",  "telemetry_snapshots", "Time-series (денормалізована)"],
            ["Аномалії",    "alert_rules, anomaly_events", "Правила + зафіксовані події"],
            ["CMMS",        "tickets, ticket_comments", "Заявки на ТО (Kanban)"],
            ["Логістика",   "missions", "Транспортні завдання"],
            ["Аудит",       "audit_log", "Журнал дій адмінів"],
        ],
        col_widths=[3.0, 7.0, 6.5])

    add_heading(doc, "4.3. Потік ingestу телеметрії", 2)
    add_paragraph(doc, "MQTT loop:")
    add_code(doc,
        "paho-mqtt thread  ─(put)─►  asyncio.Queue  ─(get)─►  consumer coroutine\n"
        "                                                           │\n"
        "                                                   session_scope()\n"
        "                                                           │\n"
        "         ┌─── TelemetrySnapshot insert ◄──── merge sections per ts ─┐\n"
        "         │                                                          │\n"
        "         ▼                                                          │\n"
        "   Robot live-state update                                          │\n"
        "   (last_x, last_y, status, last_seen_at)                           │\n"
        "         │                                                          │\n"
        "         ▼                                                          │\n"
        "   services.anomaly.evaluate_snapshot() ─► AnomalyEvent (якщо поріг)┘")

    add_heading(doc, "4.4. Предиктивна аналітика", 2)
    add_paragraph(doc,
        "Сервіс predictive обраховує 0..100 health-score для кожного "
        "компонента на основі вікна останніх ~15 хвилин телеметрії. "
        "Лінійний скалінг: value ∈ [healthy, critical] → score ∈ [100, 0]. "
        "Для батареї враховується SoH + температура; для двигуна — температура + "
        "вібрація. RUL обчислюється як (expected_life − used_hours) × health_factor, "
        "з простою моделлю впевненості 0.55 + 0.35·health_factor. Цей шар "
        "ізольований у services/predictive.py, тому його легко замінити на "
        "повноцінну ML-модель (scikit-learn / XGBoost) у майбутньому.")

    add_heading(doc, "4.5. Ролі та RBAC", 2)
    add_paragraph(doc,
        "Три ролі обрано, щоб покрити типові обов’язки цеху без вибуху "
        "привілеїв. Кожен REST-ендпоїнт захищений декларативно:")
    add_table(doc,
        ["Роль", "Обов’язки", "Ключові дозволи"],
        [
            ["admin",
             "Системний адміністратор: керує користувачами та конфігурацією; "
             "повний override поверх іншого функціоналу.",
             "USERS_MANAGE, SYSTEM_CONFIG + усі нижче"],
            ["engineer",
             "Інженер з надійності: володіє тікетами, налаштовує пороги "
             "аномалій, робить висновки з PdM-аналітики, замінює компоненти.",
             "PREDICTIVE_CONFIGURE, ALERTS_MANAGE, "
             "TICKETS_EDIT/CLOSE, ROBOTS_EDIT"],
            ["operator",
             "Диспетчер зміни: слідкує за дашбордом, створює місії, "
             "підтверджує сповіщення, вмикає базові команди (stop/resume).",
             "ROBOTS_COMMAND, MISSIONS_CREATE/CANCEL, "
             "TICKETS_CREATE, ALERTS_ACK"],
        ],
        col_widths=[2.5, 7.0, 6.5])

    add_paragraph(doc,
        "Технічно це реалізовано так: у core/roles.py описано Role (Enum) та "
        "матрицю ROLE_PERMISSIONS: Role → set[Permission]. Кожен router "
        "використовує Depends(require_permission(Permission.X)). Фронтенд "
        "повторює ту саму логіку в RequireRole-guard та у Layout-навігації.")

    add_heading(doc, "4.6. API-контракт (огляд)", 2)
    add_table(doc,
        ["Метод", "Шлях", "Дозвіл"],
        [
            ["POST", "/api/auth/login",                        "—"],
            ["POST", "/api/auth/refresh",                      "—"],
            ["GET",  "/api/auth/me",                           "authenticated"],
            ["GET",  "/api/robots",                            "ROBOTS_VIEW"],
            ["GET",  "/api/robots/{id}",                       "ROBOTS_VIEW"],
            ["POST", "/api/robots/{id}/command",               "ROBOTS_COMMAND"],
            ["GET",  "/api/telemetry/{id}/latest",             "TELEMETRY_VIEW"],
            ["GET",  "/api/telemetry/{id}/series?field=…",     "TELEMETRY_VIEW"],
            ["GET",  "/api/anomalies",                         "ALERTS_VIEW"],
            ["POST", "/api/anomalies/{id}/ack",                "ALERTS_ACK"],
            ["GET",  "/api/alert-rules",                       "ALERTS_VIEW"],
            ["POST", "/api/alert-rules",                       "ALERTS_MANAGE"],
            ["GET",  "/api/tickets",                           "TICKETS_VIEW"],
            ["POST", "/api/tickets",                           "TICKETS_CREATE"],
            ["PATCH","/api/tickets/{id}",                      "TICKETS_EDIT"],
            ["GET",  "/api/missions",                          "MISSIONS_VIEW"],
            ["POST", "/api/missions",                          "MISSIONS_CREATE"],
            ["POST", "/api/missions/{id}/cancel",              "MISSIONS_CANCEL"],
            ["GET",  "/api/predictive/fleet",                  "PREDICTIVE_VIEW"],
            ["GET",  "/api/predictive/robots/{id}/health",     "PREDICTIVE_VIEW"],
            ["GET",  "/api/predictive/robots/{id}/rul",        "PREDICTIVE_VIEW"],
            ["GET",  "/api/users",                             "USERS_MANAGE (admin)"],
            ["POST", "/api/users",                             "USERS_MANAGE (admin)"],
        ],
        col_widths=[1.8, 8.0, 5.0])

    doc.add_page_break()


    add_heading(doc, "5. Фронтенд (React + TypeScript)", 1)
    add_heading(doc, "5.1. Стек", 2)
    add_bullet(doc, "Vite 5 + React 18 + TypeScript (suspense-friendly)")
    add_bullet(doc, "TailwindCSS 3 + власні shadcn-подібні компоненти")
    add_bullet(doc, "TanStack Query 5 — кешування + refetch (раз у 2-10 с залежно від сторінки)")
    add_bullet(doc, "Zustand + persist — зберігає JWT-токени у localStorage")
    add_bullet(doc, "Recharts — графіки метрик у деталях робота")
    add_bullet(doc, "React Router 6 — маршрутизація з role-guards")

    add_heading(doc, "5.2. Навігаційна мапа", 2)
    add_code(doc,
        "/login                          (public)\n"
        "┌─ RequireAuth ──────────────────────────────────────┐\n"
        "│ /dashboard          all roles                      │\n"
        "│ /robots             all roles                      │\n"
        "│ /robots/:id         all roles (operator — view-only +\n"
        "│                                basic commands)     │\n"
        "│ /predictive         admin, engineer                │\n"
        "│ /alerts             all roles (manage — engineer+) │\n"
        "│ /tickets            all roles                      │\n"
        "│ /missions           all roles                      │\n"
        "│ /admin/users        admin                          │\n"
        "└────────────────────────────────────────────────────┘")

    add_heading(doc, "5.3. Ключові екрани", 2)
    add_table(doc,
        ["Сторінка", "Призначення", "Ролі"],
        [
            ["Dashboard", "KPI по флоту + список роботів + стрічка аномалій (refetch 5 с)", "усі"],
            ["RobotDetail", "Live-графіки (SoC, T°, вібрація), RUL-прогнози, інжекція "
                           "несправностей, команди",
             "усі (інжекція — engineer+)"],
            ["Predictive", "Рейтинг здоров’я флоту з прогрес-барами", "admin, engineer"],
            ["Alerts",     "Активні anomaly_events + CRUD alert_rules", "усі / manage — engineer+"],
            ["Tickets",    "Канбан-дошка (open / assigned / in_progress / completed) + модалка "
                           "з коментарями", "усі / edit — engineer+"],
            ["Missions",   "Список логістичних завдань + створення + скасування", "усі"],
            ["Users",      "CRUD користувачів з валідацією ролей", "admin"],
        ],
        col_widths=[3.5, 9.0, 4.0])


    add_heading(doc, "6. База даних (PostgreSQL)", 1)
    add_paragraph(doc,
        "Повна ER-схема (50+ таблиць, 4 схеми — public/telemetry/maintenance/auth) "
        "описана у файлі db.txt і доступна у діаграмах diplom_er_diagrams.html. "
        "Це еталонна модель для дипломної роботи і може бути використана як "
        "Alembic-baseline у продакшн-варіанті.")

    add_paragraph(doc,
        "У бекенді реалізовано підмножину (15 таблиць) із спрощеним ENUM ролей "
        "до трьох значень. Основні зв’язки:")
    add_code(doc,
        "                ┌───────────┐\n"
        "                │ factories │\n"
        "                └─────┬─────┘\n"
        "                      │ 1..N\n"
        "                ┌─────▼─────────────┐\n"
        "                │ production_lines  │\n"
        "                └─┬──────────┬──────┘\n"
        "                  │ 1..N     │ 1..N\n"
        "         ┌────────▼──┐   ┌───▼─────────────┐\n"
        "         │ robots    │   │ workshop_zones  │\n"
        "         └────┬──────┘   └───┬─────────────┘\n"
        "              │ 1..N         │ 1..N\n"
        "   ┌──────────▼───────┐   ┌──▼──────────────┐\n"
        "   │ robot_components │   │charging_stations│\n"
        "   └──────────┬───────┘   └─────────────────┘\n"
        "              │\n"
        "              ├─► telemetry_snapshots (N per robot, time-series)\n"
        "              │\n"
        "              ├─► anomaly_events ◄── alert_rules\n"
        "              │\n"
        "              └─► tickets (CMMS) ◄── users")


    add_heading(doc, "7. Розгортання", 1)
    add_paragraph(doc,
        "Локальне розгортання одним командою за рахунок docker-compose:")
    add_code(doc,
        "$ cd диплом\n"
        "$ docker compose up --build\n"
        "Services:\n"
        "  postgres:16   (localhost:5432,  user=amr)\n"
        "  mosquitto:2   (localhost:1883)\n"
        "  backend       (localhost:8000, /docs — Swagger)\n"
        "  frontend      (localhost:8080)\n"
        "Webots запускається окремо — відкрити webots_project/worlds/factory_floor.wbt\n"
        "і експортувати AMR_MQTT=1 у змінні середовища перед запуском симуляції.")

    add_heading(doc, "7.1. Демо-акаунти", 2)
    add_table(doc,
        ["Роль", "Email", "Пароль"],
        [
            ["admin",    "admin@progress.ua",    "admin123"],
            ["engineer", "engineer@progress.ua", "engineer123"],
            ["operator", "operator@progress.ua", "operator123"],
        ],
        col_widths=[3.0, 7.0, 5.0])


    add_heading(doc, "8. Висновки та напрями розвитку", 1)
    add_paragraph(doc,
        "Розроблена інформаційна система повноцінно демонструє повний "
        "цикл предиктивного обслуговування АМР-флоту на базі цифрових "
        "двійників: від генерації реалістичної деградаційної телеметрії "
        "у Webots до візуалізації RUL-прогнозів у веб-дашборді. "
        "Архітектура модульна — кожен рівень (simulator / broker / backend / "
        "frontend) може еволюціонувати незалежно.")
    add_paragraph(doc, "Напрями подальшого розвитку:", bold=True)
    add_bullet(doc, "Заміна евристичної оцінки RUL на ML-модель (LSTM / "
                    "XGBoost-survival) з перенавчанням на накопиченій історії.")
    add_bullet(doc, "Підключення OPC-UA шлюзу для інтеграції із "
                    "SCADA/MES/ERP (MES-replanning при низькому SoH).")
    add_bullet(doc, "Перехід InfluxDB для high-rate телеметрії (зараз все "
                    "у PostgreSQL) з Continuous Queries для downsampling.")
    add_bullet(doc, "WebSocket-канал для push-оновлень дашборда замість поточного polling.")
    add_bullet(doc, "Multi-tenant: factory_id у JWT-claims для ізоляції даних "
                    "між виробництвами.")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"[ok] saved {OUT}")


if __name__ == "__main__":
    build()
